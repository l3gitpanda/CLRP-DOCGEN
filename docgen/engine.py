"""
Core document generation engine.

Builds DOCX files from structured document data using python-docx,
with dark-theme styling, gold heading bars, and Republic emblem
matching the Clone Wars RP document format.
"""

import os
import subprocess
import shutil
from datetime import datetime
from typing import List, Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from .styles import StyleConfig, COLORS


# Republic cog / gear emblem drawn with Unicode
REPUBLIC_EMBLEM = "\u2699"  # ⚙  gear symbol


class DocumentEngine:
    """Generates polished DOCX (and optionally PDF) training documents."""

    def __init__(self, style: StyleConfig):
        self.style = style
        self.doc = Document()
        self._setup_page_layout()
        self._toc_entries: List[str] = []

    # ------------------------------------------------------------------
    # Page layout & dark background
    # ------------------------------------------------------------------

    def _setup_page_layout(self):
        """Configure page size, margins, dark background, and defaults."""
        for section in self.doc.sections:
            section.page_width = Inches(self.style.page_width)
            section.page_height = Inches(self.style.page_height)
            section.top_margin = Inches(self.style.margin_top)
            section.bottom_margin = Inches(self.style.margin_bottom)
            section.left_margin = Inches(self.style.margin_left)
            section.right_margin = Inches(self.style.margin_right)

        # Set document background color
        if self.style.dark_mode:
            bg_hex = self.style._bg_hex()
            background = parse_xml(
                f'<w:background {nsdecls("w")} w:color="{bg_hex}"/>'
            )
            self.doc.element.insert(0, background)
            # Also need displayBackgroundShape in settings
            settings = self.doc.settings.element
            disp_bg = parse_xml(
                f'<w:displayBackgroundShape {nsdecls("w")}/>'
            )
            settings.append(disp_bg)

        # Set default paragraph style
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = self.style.body_font
        font.size = Pt(self.style.body_size)
        font.color.rgb = self.style.resolve_color(self.style.body_color)
        pf = style.paragraph_format
        pf.space_before = Pt(self.style.paragraph_spacing_before)
        pf.space_after = Pt(self.style.paragraph_spacing_after)
        pf.line_spacing = self.style.line_spacing

    def _shade_paragraph(self, paragraph):
        """Apply page background shading to a paragraph for compatibility."""
        if self.style.dark_mode:
            bg_hex = self.style._bg_hex()
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{bg_hex}" w:val="clear"/>'
            )
            paragraph._p.get_or_add_pPr().append(shading)

    # ------------------------------------------------------------------
    # Helper: apply font to a run
    # ------------------------------------------------------------------

    def _apply_run_style(self, run, font_name=None, size=None,
                         color_key=None, bold=None, italic=None,
                         underline=None):
        """Apply formatting to a single run."""
        if font_name:
            run.font.name = font_name
        if size:
            run.font.size = Pt(size)
        if color_key:
            run.font.color.rgb = self.style.resolve_color(color_key)
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        if underline is not None:
            run.underline = underline

    # ------------------------------------------------------------------
    # Heading bar (gold bar with dark text)
    # ------------------------------------------------------------------

    def _add_heading_bar(self, text: str, level: int = 1):
        """Add a full-width gold heading bar using a shaded table cell."""
        bar_hex = self.style._heading_bar_hex()
        bg_hex = self.style._bg_hex()

        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Make table span full width
        tbl = table._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
            f'<w:tblPr {nsdecls("w")}/>'
        )
        # Set table width to 100%
        tbl_w = parse_xml(
            f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'
        )
        # Remove existing tblW if present
        for existing in tbl_pr.findall(qn('w:tblW')):
            tbl_pr.remove(existing)
        tbl_pr.append(tbl_w)

        # Remove table borders (the bar is the fill itself)
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:color="{bar_hex}"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:color="{bar_hex}"/>'
            f'  <w:left w:val="single" w:sz="4" w:color="{bar_hex}"/>'
            f'  <w:right w:val="single" w:sz="4" w:color="{bar_hex}"/>'
            f'  <w:insideH w:val="none" w:sz="0" w:color="000000"/>'
            f'  <w:insideV w:val="none" w:sz="0" w:color="000000"/>'
            f'</w:tblBorders>'
        )
        for existing in tbl_pr.findall(qn('w:tblBorders')):
            tbl_pr.remove(existing)
        tbl_pr.append(borders)

        cell = table.cell(0, 0)
        cell.text = ""

        # Apply gold background fill
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{bar_hex}" w:val="clear"/>'
        )
        tc_pr.append(shading)

        # Cell margins for padding
        tc_mar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'  <w:top w:w="60" w:type="dxa"/>'
            f'  <w:bottom w:w="60" w:type="dxa"/>'
            f'  <w:left w:w="115" w:type="dxa"/>'
            f'  <w:right w:w="115" w:type="dxa"/>'
            f'</w:tcMar>'
        )
        tc_pr.append(tc_mar)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

        if level == 1:
            run = p.add_run(text)
            self._apply_run_style(
                run,
                font_name=self.style.heading_font,
                size=self.style.heading_size,
                color_key=self.style.heading_bar_text_color,
                bold=self.style.heading_bold,
            )
        else:
            run = p.add_run(text)
            self._apply_run_style(
                run,
                font_name=self.style.subheading_font,
                size=self.style.subheading_size - 1,
                color_key=self.style.heading_bar_text_color,
                bold=True,
            )

    # ------------------------------------------------------------------
    # Title page
    # ------------------------------------------------------------------

    def add_title_page(self, title: str, subtitle: str = "",
                       author: str = "", formatted_by: str = "",
                       version_date: str = "", extra_lines: List[str] = None):
        """Create a formatted title page with emblem and dark background."""
        # Republic emblem
        if self.style.show_emblem:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._shade_paragraph(p)
            run = p.add_run(REPUBLIC_EMBLEM)
            self._apply_run_style(
                run,
                size=72,
                color_key=self.style.accent_color,
            )

        # Spacing
        for _ in range(2):
            p = self.doc.add_paragraph("")
            self._shade_paragraph(p)

        # Title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._shade_paragraph(p)
        run = p.add_run(title)
        self._apply_run_style(
            run,
            font_name=self.style.title_font,
            size=self.style.title_size,
            color_key=self.style.title_color,
            bold=self.style.title_bold,
        )

        # Subtitle
        if subtitle:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._shade_paragraph(p)
            run = p.add_run(subtitle)
            self._apply_run_style(
                run,
                font_name=self.style.subtitle_font,
                size=self.style.subtitle_size,
                color_key=self.style.subtitle_color,
            )

        # Gold divider line
        self._add_gold_line()

        # Version / date
        if version_date:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._shade_paragraph(p)
            run = p.add_run(f"Latest Version: {version_date}")
            self._apply_run_style(
                run,
                font_name=self.style.subtitle_font,
                size=self.style.subtitle_size - 2,
                color_key=self.style.subtitle_color,
                italic=True,
            )

        # Author / formatted by
        if author or formatted_by:
            p = self.doc.add_paragraph("")
            self._shade_paragraph(p)
            if author:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._shade_paragraph(p)
                run = p.add_run(f"Created by: {author}")
                self._apply_run_style(
                    run, size=10, color_key="medium_gray", italic=True
                )
            if formatted_by:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._shade_paragraph(p)
                run = p.add_run(f"Formatted by: {formatted_by}")
                self._apply_run_style(
                    run, size=10, color_key="medium_gray", italic=True
                )

        # Extra lines
        if extra_lines:
            p = self.doc.add_paragraph("")
            self._shade_paragraph(p)
            for line in extra_lines:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._shade_paragraph(p)
                run = p.add_run(line)
                self._apply_run_style(run, size=10, color_key="medium_gray")

        # Page break
        self.doc.add_page_break()

    # ------------------------------------------------------------------
    # Table of contents
    # ------------------------------------------------------------------

    def add_table_of_contents(self, entries: List[dict] = None):
        """Add a Table of Contents page with gold bullets."""
        self._add_heading_bar(self.style.toc_title)
        p = self.doc.add_paragraph("")
        self._shade_paragraph(p)

        items = entries or [{"title": t} for t in self._toc_entries]
        symbol = self.style.section_symbol or "\u2022"  # bullet

        for item in items:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            self._shade_paragraph(p)

            # Gold bullet
            bullet_run = p.add_run(f"{symbol}  ")
            self._apply_run_style(
                bullet_run,
                font_name=self.style.body_font,
                size=self.style.body_size,
                color_key=self.style.accent_color,
                bold=True,
            )

            # Entry text
            title_text = item.get("title", "")
            text_run = p.add_run(title_text)
            self._apply_run_style(
                text_run,
                font_name=self.style.body_font,
                size=self.style.body_size,
                color_key=self.style.body_color,
            )

        self.doc.add_page_break()

    # ------------------------------------------------------------------
    # Headings
    # ------------------------------------------------------------------

    def add_heading(self, text: str, level: int = 1, track_toc: bool = True):
        """Add a styled heading. level 1 = gold bar, level 2 = gold text."""
        if track_toc and level <= 2:
            self._toc_entries.append(text)

        if self.style.use_section_symbols and self.style.section_symbol:
            text = f"{self.style.section_symbol} {text} {self.style.section_symbol}"

        if level == 1:
            # Add small spacer before the bar
            p = self.doc.add_paragraph("")
            self._shade_paragraph(p)
            self._add_heading_bar(text, level=1)
        else:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            self._shade_paragraph(p)
            run = p.add_run(text)
            self._apply_run_style(
                run,
                font_name=self.style.subheading_font,
                size=self.style.subheading_size,
                color_key=self.style.subheading_color,
                bold=self.style.subheading_bold,
            )

    # ------------------------------------------------------------------
    # Body text
    # ------------------------------------------------------------------

    def add_paragraph(self, text: str, color_key: str = None,
                      bold: bool = False, italic: bool = False,
                      alignment: str = "left", indent: float = 0):
        """Add a styled body paragraph."""
        p = self.doc.add_paragraph()

        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        p.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self._shade_paragraph(p)

        if indent > 0:
            p.paragraph_format.left_indent = Inches(indent)

        run = p.add_run(text)
        self._apply_run_style(
            run,
            font_name=self.style.body_font,
            size=self.style.body_size,
            color_key=color_key or self.style.body_color,
            bold=bold or self.style.body_bold,
            italic=italic,
        )
        return p

    def add_colored_text(self, segments: list):
        """Add a paragraph with mixed-color segments.

        segments: list of (text, color_key, bold, italic) tuples.
        """
        p = self.doc.add_paragraph()
        self._shade_paragraph(p)
        for seg in segments:
            text = seg[0]
            color = seg[1] if len(seg) > 1 else None
            bold = seg[2] if len(seg) > 2 else False
            italic = seg[3] if len(seg) > 3 else False
            run = p.add_run(text)
            self._apply_run_style(
                run,
                font_name=self.style.body_font,
                size=self.style.body_size,
                color_key=color or self.style.body_color,
                bold=bold,
                italic=italic,
            )
        return p

    # ------------------------------------------------------------------
    # Instructional text (color-coded)
    # ------------------------------------------------------------------

    def add_read_aloud(self, text: str):
        """Add text meant to be read aloud to recruits (green)."""
        return self.add_paragraph(
            text, color_key=self.style.read_aloud_color, italic=True
        )

    def add_host_info(self, text: str):
        """Add host-only information (red)."""
        return self.add_paragraph(
            text, color_key=self.style.host_info_color, bold=True
        )

    def add_important_info(self, text: str):
        """Add important info (blue)."""
        return self.add_paragraph(
            text, color_key=self.style.important_info_color, bold=True
        )

    # ------------------------------------------------------------------
    # Lists
    # ------------------------------------------------------------------

    def add_bullet_list(self, items: List[str], indent: float = 0.4,
                        color_key: str = None):
        """Add a bulleted list with gold bullets."""
        for item in items:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(indent)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            self._shade_paragraph(p)

            # Gold bullet marker
            bullet_run = p.add_run("\u2022  ")
            self._apply_run_style(
                bullet_run,
                font_name=self.style.body_font,
                size=self.style.body_size,
                color_key=self.style.accent_color,
                bold=True,
            )

            run = p.add_run(item)
            self._apply_run_style(
                run,
                font_name=self.style.body_font,
                size=self.style.body_size,
                color_key=color_key or self.style.body_color,
            )

    def add_numbered_list(self, items: List[str], indent: float = 0.4,
                          color_key: str = None, start_num: int = 1):
        """Add a numbered list with gold numbers."""
        for i, item in enumerate(items, start=start_num):
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(indent)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            self._shade_paragraph(p)

            # Gold number
            num_run = p.add_run(f"{i}. ")
            self._apply_run_style(
                num_run,
                font_name=self.style.body_font,
                size=self.style.body_size,
                color_key=self.style.accent_color,
                bold=True,
            )

            run = p.add_run(item)
            self._apply_run_style(
                run,
                font_name=self.style.body_font,
                size=self.style.body_size,
                color_key=color_key or self.style.body_color,
            )

    def add_lettered_sub_list(self, items: List[str], indent: float = 0.7,
                              color_key: str = None):
        """Add a sub-list with letter labels (a., b., c., ...)."""
        for i, item in enumerate(items):
            letter = chr(ord("a") + i)
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(indent)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            self._shade_paragraph(p)

            ltr_run = p.add_run(f"{letter}. ")
            self._apply_run_style(
                ltr_run,
                font_name=self.style.body_font,
                size=self.style.body_size,
                color_key=self.style.accent_color,
                bold=True,
            )

            run = p.add_run(item)
            self._apply_run_style(
                run,
                font_name=self.style.body_font,
                size=self.style.body_size,
                color_key=color_key or self.style.body_color,
            )

    # ------------------------------------------------------------------
    # Q&A blocks
    # ------------------------------------------------------------------

    def add_qa_block(self, question: str, answer: str, q_label: str = "Q",
                     a_label: str = "A"):
        """Add a styled question/answer pair."""
        # Question
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        self._shade_paragraph(p)
        q_run = p.add_run(f"{q_label}: {question}")
        self._apply_run_style(
            q_run,
            font_name=self.style.body_font,
            size=self.style.body_size,
            color_key=self.style.accent_color,
            bold=True,
        )

        # Answer
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.6)
        self._shade_paragraph(p)
        a_run = p.add_run(f"{a_label}: {answer}")
        self._apply_run_style(
            a_run,
            font_name=self.style.body_font,
            size=self.style.body_size,
            color_key=self.style.body_color,
            italic=True,
        )

    # ------------------------------------------------------------------
    # Tables
    # ------------------------------------------------------------------

    def add_table(self, headers: List[str], rows: List[List[str]],
                  col_widths: List[float] = None):
        """Add a styled table with gold header row on dark background."""
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        bar_hex = self.style._heading_bar_hex()
        bg_hex = self.style._bg_hex()
        accent = self.style.resolve_color(self.style.accent_color)
        accent_hex = f"{accent[0]:02X}{accent[1]:02X}{accent[2]:02X}"

        # Table borders
        tbl = table._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
            f'<w:tblPr {nsdecls("w")}/>'
        )
        tbl_borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:color="{accent_hex}"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:color="{accent_hex}"/>'
            f'  <w:left w:val="single" w:sz="4" w:color="{accent_hex}"/>'
            f'  <w:right w:val="single" w:sz="4" w:color="{accent_hex}"/>'
            f'  <w:insideH w:val="single" w:sz="4" w:color="{accent_hex}"/>'
            f'  <w:insideV w:val="single" w:sz="4" w:color="{accent_hex}"/>'
            f'</w:tblBorders>'
        )
        for existing in tbl_pr.findall(qn('w:tblBorders')):
            tbl_pr.remove(existing)
        tbl_pr.append(tbl_borders)

        # Header row
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = ""
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header)
            self._apply_run_style(
                run,
                font_name=self.style.heading_font,
                size=self.style.body_size,
                color_key=self.style.heading_bar_text_color,
                bold=True,
            )
            # Gold header background
            tc_pr = hdr_cells[i]._tc.get_or_add_tcPr()
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{bar_hex}" w:val="clear"/>'
            )
            tc_pr.append(shading)

        # Data rows
        for r_idx, row in enumerate(rows):
            cells = table.rows[r_idx + 1].cells
            for c_idx, cell_text in enumerate(row):
                cells[c_idx].text = ""
                p = cells[c_idx].paragraphs[0]
                run = p.add_run(str(cell_text))
                self._apply_run_style(
                    run,
                    font_name=self.style.body_font,
                    size=self.style.body_size - 1,
                    color_key=self.style.body_color,
                )
                # Dark cell background
                if self.style.dark_mode:
                    tc_pr = cells[c_idx]._tc.get_or_add_tcPr()
                    shading = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="{bg_hex}" w:val="clear"/>'
                    )
                    tc_pr.append(shading)

        # Column widths
        if col_widths:
            for row in table.rows:
                for i, width in enumerate(col_widths):
                    if i < len(row.cells):
                        row.cells[i].width = Inches(width)

        return table

    # ------------------------------------------------------------------
    # Chain of command
    # ------------------------------------------------------------------

    def add_chain_of_command(self, chain: List[str]):
        """Add a vertical chain of command with arrow connectors."""
        for i, rank in enumerate(chain):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._shade_paragraph(p)
            run = p.add_run(rank)
            self._apply_run_style(
                run,
                font_name=self.style.heading_font,
                size=self.style.body_size + 1,
                color_key=self.style.accent_color,
                bold=True,
            )

            if i < len(chain) - 1:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                self._shade_paragraph(p)
                run = p.add_run("\u2193")  # down arrow
                self._apply_run_style(
                    run,
                    size=self.style.body_size + 6,
                    color_key=self.style.accent_color,
                    bold=True,
                )

    # ------------------------------------------------------------------
    # Visual elements
    # ------------------------------------------------------------------

    def _add_gold_line(self):
        """Add a thin gold horizontal line."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        self._shade_paragraph(p)

        divider_char = "\u2500"  # box-drawing horizontal
        run = p.add_run(divider_char * 60)
        self._apply_run_style(
            run,
            size=6,
            color_key=self.style.divider_color,
        )

    def _add_divider(self):
        """Add a gold divider (alias)."""
        self._add_gold_line()

    def add_divider(self):
        """Public method to add a divider."""
        self._add_gold_line()

    def add_spacer(self, lines: int = 1):
        """Add blank lines for spacing."""
        for _ in range(lines):
            p = self.doc.add_paragraph("")
            self._shade_paragraph(p)

    def add_page_break(self):
        """Add a page break."""
        self.doc.add_page_break()

    # ------------------------------------------------------------------
    # Color code legend
    # ------------------------------------------------------------------

    def add_color_code_legend(self):
        """Add a color code explanation block."""
        self._add_heading_bar("Color Codes", level=2)

        codes = [
            ("Green text", self.style.read_aloud_color,
             "  — you will read aloud to the trainee."),
            ("Red text", self.style.host_info_color,
             "  — is information that you will need; Do not read aloud."),
            ("Blue text", self.style.important_info_color,
             "  — is important information (adverts, droid numbers, etc.)"),
        ]

        for label, color, desc in codes:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            self._shade_paragraph(p)
            colored_run = p.add_run(label)
            self._apply_run_style(
                colored_run,
                font_name=self.style.body_font,
                size=self.style.body_size,
                color_key=color,
                bold=True,
            )
            desc_run = p.add_run(desc)
            self._apply_run_style(
                desc_run,
                font_name=self.style.body_font,
                size=self.style.body_size,
                color_key=self.style.body_color,
            )

    # ------------------------------------------------------------------
    # Info box / callout
    # ------------------------------------------------------------------

    def add_callout_box(self, text: str, style_type: str = "info"):
        """Add a bordered callout box on dark background."""
        color_map = {
            "info": self.style.important_info_color,
            "warning": self.style.host_info_color,
            "note": self.style.accent_color,
        }
        color = color_map.get(style_type, self.style.accent_color)

        bg_hex = self.style._bg_hex()
        resolved = self.style.resolve_color(color)
        border_hex = f"{resolved[0]:02X}{resolved[1]:02X}{resolved[2]:02X}"

        # Slightly lighter dark background for callout
        callout_bg = self.style.resolve_color("callout_bg")
        callout_hex = f"{callout_bg[0]:02X}{callout_bg[1]:02X}{callout_bg[2]:02X}"

        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Full-width
        tbl_pr = table._tbl.tblPr
        tbl_w = parse_xml(
            f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'
        )
        for existing in tbl_pr.findall(qn('w:tblW')):
            tbl_pr.remove(existing)
        tbl_pr.append(tbl_w)

        cell = table.cell(0, 0)
        cell.text = ""

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        self._apply_run_style(
            run,
            font_name=self.style.body_font,
            size=self.style.body_size,
            color_key=color,
            bold=True,
        )

        # Border and background
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="8" w:color="{border_hex}"/>'
            f'  <w:bottom w:val="single" w:sz="8" w:color="{border_hex}"/>'
            f'  <w:left w:val="single" w:sz="8" w:color="{border_hex}"/>'
            f'  <w:right w:val="single" w:sz="8" w:color="{border_hex}"/>'
            f'</w:tcBorders>'
        )
        tc_pr.append(borders)

        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{callout_hex}" w:val="clear"/>'
        )
        tc_pr.append(shading)

        p = self.doc.add_paragraph("")
        self._shade_paragraph(p)

    # ------------------------------------------------------------------
    # Header / footer metadata line
    # ------------------------------------------------------------------

    def add_metadata_line(self, author: str = "", formatted_by: str = "",
                          created: str = "", updated: str = "",
                          alignment: str = "right"):
        """Add a small metadata attribution line."""
        parts = []
        if author:
            parts.append(f"Created By: {author}")
        if formatted_by:
            parts.append(f"Formatted by: {formatted_by}")
        if created:
            parts.append(f"Created: {created}")
        if updated:
            parts.append(f"Updated: {updated}")

        if not parts:
            return

        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }

        p = self.doc.add_paragraph()
        p.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.RIGHT)
        self._shade_paragraph(p)
        text = " | ".join(parts)
        run = p.add_run(text)
        self._apply_run_style(
            run,
            size=8,
            color_key="medium_gray",
            italic=True,
        )

    # ------------------------------------------------------------------
    # Save
    # ------------------------------------------------------------------

    def save_docx(self, filepath: str) -> str:
        """Save the document as .docx and return the path."""
        os.makedirs(os.path.dirname(filepath) or ".", exist_ok=True)
        self.doc.save(filepath)
        return filepath

    def save_pdf(self, filepath: str) -> str:
        """Save the document as PDF.

        Tries LibreOffice first for best fidelity, then falls back to
        a basic text-based PDF.
        """
        # First save as docx
        docx_path = filepath.rsplit(".", 1)[0] + ".docx"
        self.save_docx(docx_path)

        # Try LibreOffice conversion
        lo_path = shutil.which("libreoffice") or shutil.which("soffice")
        if lo_path:
            out_dir = os.path.dirname(filepath) or "."
            result = subprocess.run(
                [lo_path, "--headless", "--convert-to", "pdf",
                 "--outdir", out_dir, docx_path],
                capture_output=True, timeout=60,
            )
            expected_pdf = docx_path.rsplit(".", 1)[0] + ".pdf"
            if os.path.exists(expected_pdf):
                if expected_pdf != filepath:
                    os.rename(expected_pdf, filepath)
                return filepath

        # Fallback: minimal PDF
        return self._minimal_pdf_fallback(filepath)

    def _minimal_pdf_fallback(self, filepath: str) -> str:
        """Generate a basic PDF from the document's text content."""
        lines = []
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if text:
                is_heading = False
                for run in para.runs:
                    if run.font.size and run.font.size >= Pt(14):
                        is_heading = True
                        break
                lines.append((text, is_heading))
            else:
                lines.append(("", False))

        objects = []
        obj_id = 0

        def add_obj(content):
            nonlocal obj_id
            obj_id += 1
            objects.append((obj_id, content))
            return obj_id

        catalog_id = add_obj(None)
        pages_id = add_obj(None)
        font_id = add_obj(
            b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"
        )
        font_bold_id = add_obj(
            b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold >>"
        )

        page_width, page_height = 612, 792
        margin = 72
        y = page_height - margin
        content_streams = []
        current_stream_lines = []

        def flush_page():
            nonlocal y
            stream_text = "\n".join(current_stream_lines)
            stream_bytes = stream_text.encode("latin-1", errors="replace")
            current_stream_lines.clear()
            content_streams.append(stream_bytes)
            y = page_height - margin

        def wrap_text(text, chars_per_line=85):
            words = text.split()
            result_lines = []
            current = ""
            for word in words:
                if len(current) + len(word) + 1 <= chars_per_line:
                    current = f"{current} {word}" if current else word
                else:
                    if current:
                        result_lines.append(current)
                    current = word
            if current:
                result_lines.append(current)
            return result_lines if result_lines else [""]

        def escape_pdf(text):
            return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

        # Dark background rectangle
        current_stream_lines.append(
            f"0.231 0.231 0.231 rg "
            f"0 0 {page_width} {page_height} re f"
        )
        current_stream_lines.append("BT")
        for text, is_heading in lines:
            if not text:
                y -= 12
                if y < margin:
                    current_stream_lines.append("ET")
                    flush_page()
                    current_stream_lines.append(
                        f"0.231 0.231 0.231 rg "
                        f"0 0 {page_width} {page_height} re f"
                    )
                    current_stream_lines.append("BT")
                continue

            font_size = 16 if is_heading else 11
            line_height = font_size * 1.4
            font_ref = "/F2" if is_heading else "/F1"

            # Set text color to white for dark bg
            current_stream_lines.append("0.91 0.91 0.91 rg")

            wrapped = wrap_text(text, 75 if is_heading else 85)
            for wline in wrapped:
                if y - line_height < margin:
                    current_stream_lines.append("ET")
                    flush_page()
                    current_stream_lines.append(
                        f"0.231 0.231 0.231 rg "
                        f"0 0 {page_width} {page_height} re f"
                    )
                    current_stream_lines.append("BT")
                    current_stream_lines.append("0.91 0.91 0.91 rg")

                safe_text = escape_pdf(wline)
                x = margin
                if is_heading:
                    approx_width = len(wline) * font_size * 0.5
                    x = max(margin, (page_width - approx_width) / 2)

                current_stream_lines.append(
                    f"{font_ref} {font_size} Tf "
                    f"{x:.0f} {y:.0f} Td "
                    f"({safe_text}) Tj"
                )
                y -= line_height

        current_stream_lines.append("ET")
        flush_page()

        page_ids = []
        for stream_bytes in content_streams:
            stream_id = add_obj(stream_bytes)
            page_id = add_obj(None)
            page_ids.append((page_id, stream_id))

        pdf_parts = [b"%PDF-1.4\n"]
        offsets = {}

        def write_obj(oid, data):
            offsets[oid] = len(b"".join(pdf_parts))
            pdf_parts.append(f"{oid} 0 obj\n".encode())
            pdf_parts.append(data)
            pdf_parts.append(b"\nendobj\n")

        write_obj(catalog_id,
                  f"<< /Type /Catalog /Pages {pages_id} 0 R >>".encode())
        kids = " ".join(f"{pid} 0 R" for pid, _ in page_ids)
        write_obj(pages_id,
                  f"<< /Type /Pages /Kids [{kids}] /Count {len(page_ids)} >>".encode())
        write_obj(font_id,
                  b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
        write_obj(font_bold_id,
                  b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold >>")

        for page_id, stream_id in page_ids:
            stream_data = None
            for oid, content in objects:
                if oid == stream_id:
                    stream_data = content
                    break

            write_obj(stream_id,
                      f"<< /Length {len(stream_data)} >>\nstream\n".encode()
                      + stream_data
                      + b"\nendstream")

            write_obj(page_id,
                      (f"<< /Type /Page /Parent {pages_id} 0 R "
                       f"/MediaBox [0 0 {page_width} {page_height}] "
                       f"/Contents {stream_id} 0 R "
                       f"/Resources << /Font << /F1 {font_id} 0 R "
                       f"/F2 {font_bold_id} 0 R >> >> >>").encode())

        xref_offset = len(b"".join(pdf_parts))
        pdf_parts.append(b"xref\n")
        pdf_parts.append(f"0 {obj_id + 1}\n".encode())
        pdf_parts.append(b"0000000000 65535 f \n")
        for oid in range(1, obj_id + 1):
            offset = offsets.get(oid, 0)
            pdf_parts.append(f"{offset:010d} 00000 n \n".encode())

        pdf_parts.append(b"trailer\n")
        pdf_parts.append(
            f"<< /Size {obj_id + 1} /Root {catalog_id} 0 R >>\n".encode()
        )
        pdf_parts.append(b"startxref\n")
        pdf_parts.append(f"{xref_offset}\n".encode())
        pdf_parts.append(b"%%EOF\n")

        os.makedirs(os.path.dirname(filepath) or ".", exist_ok=True)
        with open(filepath, "wb") as f:
            f.write(b"".join(pdf_parts))

        return filepath

    def save(self, filepath: str, fmt: str = "docx") -> str:
        """Save in the requested format. fmt: 'docx' or 'pdf'."""
        if fmt.lower() == "pdf":
            return self.save_pdf(filepath)
        return self.save_docx(filepath)
